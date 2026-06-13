"""
===============================================================================
  Họ và tên   : Huỳnh Phát Lợi
  MSHV        : KHMT836016
  Môn         : Computer Vision
  Bài tập ngày: 06/06/2026
===============================================================================

Face recognition (face / non-face classification) using LBPH features + SVM.

Assignment
----------
Use LBPH (Local Binary Pattern Histogram) feature extraction together with an
SVM classifier (libsvm) to build a face recognition program.

  * Dataset : Kaggle "fareselmenshawii/face-detection-dataset"
              (images + YOLO-format bounding boxes for faces)
  * SVM     : libsvm. scikit-learn's `SVC` is a thin Python wrapper around
              libsvm (https://www.csie.ntu.edu.tw/~cjlin/libsvm/), so we use it.

Because the dataset ships *detection* labels (face bounding boxes) instead of
*identity* labels, the recognition task here is framed as a binary classifier:

        FACE  (label 1)  -> image patches cropped from the bounding boxes
        NON-FACE (label 0) -> random background patches that do not overlap faces

The model is trained on the `train` split and evaluated on the `val` split.

Expected dataset layout (auto-detected, the script is tolerant to variations):

    exam3/dataset/
        images/train/*.jpg   labels/train/*.txt
        images/val/*.jpg     labels/val/*.txt

Each YOLO label line:  <class> <x_center> <y_center> <width> <height>
with all coordinates normalised to [0, 1].

Run:
    python3 face_recognition_lbph_svm.py
"""

# =============================================================================
# Step 0 - Imports and global configuration
# =============================================================================
import os
import glob
import random
import datetime

import cv2
import numpy as np
import matplotlib
matplotlib.use("Agg")            # headless backend: save figures to files
import matplotlib.pyplot as plt
from skimage.feature import local_binary_pattern
from sklearn.svm import SVC
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    classification_report,
    confusion_matrix,
)

# --- Student / assignment information ----------------------------------------
STUDENT_NAME = "Huỳnh Phát Lợi"
STUDENT_ID = "KHMT836016"
COURSE = "Computer Vision"
ASSIGNMENT_DATE = "06/06/2026"

# --- Reproducibility ---------------------------------------------------------
RANDOM_SEED = 42
random.seed(RANDOM_SEED)
np.random.seed(RANDOM_SEED)

# --- Paths -------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATASET_DIR = os.path.join(SCRIPT_DIR, "dataset")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "outputs")   # charts + Word report go here

# --- Patch / feature parameters ----------------------------------------------
PATCH_SIZE = (64, 64)        # every face / non-face patch is resized to this
LBP_POINTS = 8               # number of circularly symmetric neighbour points
LBP_RADIUS = 1               # radius of the LBP circle
LBP_METHOD = "uniform"       # "uniform" -> P + 2 distinct patterns
GRID = (8, 8)                # patch is split into GRID cells; one histogram/cell

# --- Sampling limits (keep the run fast; raise for higher quality) -----------
MAX_IMAGES_TRAIN = None      # None -> use ALL train images (full dataset)
MAX_IMAGES_VAL = None        # None -> use ALL val images (full dataset)
MAX_FACES_PER_IMAGE = 3      # cap positives taken from a single image
NEG_PER_IMAGE = 3            # random background patches sampled per image
MIN_BOX_SIZE = 20            # ignore tiny boxes (too small to be informative)


# =============================================================================
# Step 1 - Locate the dataset directories (robust auto-detection)
# =============================================================================
def find_split_dirs(dataset_dir):
    """Return {'train': (images_dir, labels_dir), 'val': (...)}.

    The dataset may be nested in an extra sub-folder (e.g.
    dataset/face-detection-dataset/images/train), so we search recursively for
    an `images/<split>` folder that has a matching `labels/<split>` sibling.
    Handled layouts: images/<split> + labels/<split>, and <split>/images +
    <split>/labels. The validation split may be named "val" or "valid".
    """
    split_aliases = {"train": ["train"], "val": ["val", "valid"]}
    splits = {}
    for root, dirs, _ in os.walk(dataset_dir):
        for split, names in split_aliases.items():
            if split in splits:
                continue
            for name in names:
                # layout A: <root>/images/<name> + <root>/labels/<name>
                img_a = os.path.join(root, "images", name)
                lbl_a = os.path.join(root, "labels", name)
                if os.path.isdir(img_a) and os.path.isdir(lbl_a):
                    splits[split] = (img_a, lbl_a)
                    break
                # layout B: <root>/<name>/images + <root>/<name>/labels
                img_b = os.path.join(root, name, "images")
                lbl_b = os.path.join(root, name, "labels")
                if os.path.isdir(img_b) and os.path.isdir(lbl_b):
                    splits[split] = (img_b, lbl_b)
                    break
        if "train" in splits and "val" in splits:
            break
    return splits


def list_images(img_dir):
    """Return a sorted list of image file paths inside `img_dir`."""
    exts = ("*.jpg", "*.jpeg", "*.png", "*.bmp")
    files = []
    for ext in exts:
        files.extend(glob.glob(os.path.join(img_dir, ext)))
        files.extend(glob.glob(os.path.join(img_dir, ext.upper())))
    return sorted(set(files))


# =============================================================================
# Step 2 - Read YOLO labels and crop face / non-face patches
# =============================================================================
def read_yolo_labels(label_path, img_w, img_h):
    """Parse a YOLO .txt file into a list of pixel boxes [x1, y1, x2, y2]."""
    boxes = []
    if not os.path.isfile(label_path):
        return boxes
    with open(label_path, "r") as fh:
        for line in fh:
            parts = line.split()
            if len(parts) < 5:
                continue
            # parts[0] is the class id (ignored: single "face" class)
            xc, yc, bw, bh = map(float, parts[1:5])
            x1 = int((xc - bw / 2.0) * img_w)
            y1 = int((yc - bh / 2.0) * img_h)
            x2 = int((xc + bw / 2.0) * img_w)
            y2 = int((yc + bh / 2.0) * img_h)
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(img_w, x2), min(img_h, y2)
            if x2 - x1 >= MIN_BOX_SIZE and y2 - y1 >= MIN_BOX_SIZE:
                boxes.append([x1, y1, x2, y2])
    return boxes


def iou(box_a, box_b):
    """Intersection-over-Union of two [x1, y1, x2, y2] boxes."""
    xa = max(box_a[0], box_b[0])
    ya = max(box_a[1], box_b[1])
    xb = min(box_a[2], box_b[2])
    yb = min(box_a[3], box_b[3])
    inter = max(0, xb - xa) * max(0, yb - ya)
    if inter == 0:
        return 0.0
    area_a = (box_a[2] - box_a[0]) * (box_a[3] - box_a[1])
    area_b = (box_b[2] - box_b[0]) * (box_b[3] - box_b[1])
    return inter / float(area_a + area_b - inter)


def sample_negative_boxes(face_boxes, img_w, img_h, n_samples):
    """Sample random boxes that barely overlap any face box (background)."""
    negatives = []
    if not face_boxes:
        # no faces in this image -> every region is background
        ref_sizes = [(60, 60)]
    else:
        ref_sizes = [(b[2] - b[0], b[3] - b[1]) for b in face_boxes]
    attempts = 0
    while len(negatives) < n_samples and attempts < n_samples * 25:
        attempts += 1
        bw, bh = random.choice(ref_sizes)
        bw = int(np.clip(bw, MIN_BOX_SIZE, img_w))
        bh = int(np.clip(bh, MIN_BOX_SIZE, img_h))
        if bw >= img_w or bh >= img_h:
            continue
        x1 = random.randint(0, img_w - bw)
        y1 = random.randint(0, img_h - bh)
        cand = [x1, y1, x1 + bw, y1 + bh]
        if all(iou(cand, fb) < 0.1 for fb in face_boxes):
            negatives.append(cand)
    return negatives


def crop_patch(gray_img, box):
    """Crop `box` from a grayscale image and resize to PATCH_SIZE."""
    x1, y1, x2, y2 = box
    patch = gray_img[y1:y2, x1:x2]
    if patch.size == 0:
        return None
    return cv2.resize(patch, PATCH_SIZE, interpolation=cv2.INTER_AREA)


# =============================================================================
# Step 3 - LBPH feature extraction
# =============================================================================
def extract_lbph(patch):
    """Compute the LBPH descriptor of a grayscale patch.

    The patch is divided into a GRID of cells. For every cell we build a
    histogram of (uniform) LBP codes; concatenating and L1-normalising those
    histograms yields the final feature vector -- this is exactly the "Local
    Binary Pattern *Histogram*" representation.
    """
    lbp = local_binary_pattern(patch, LBP_POINTS, LBP_RADIUS, LBP_METHOD)
    n_bins = LBP_POINTS + 2  # number of distinct codes for the "uniform" method

    cell_h = patch.shape[0] // GRID[0]
    cell_w = patch.shape[1] // GRID[1]

    features = []
    for gy in range(GRID[0]):
        for gx in range(GRID[1]):
            cell = lbp[gy * cell_h:(gy + 1) * cell_h,
                       gx * cell_w:(gx + 1) * cell_w]
            hist, _ = np.histogram(cell.ravel(), bins=n_bins,
                                   range=(0, n_bins))
            hist = hist.astype(np.float32)
            hist /= (hist.sum() + 1e-6)   # L1 normalise -> illumination robust
            features.extend(hist)
    return np.asarray(features, dtype=np.float32)


# =============================================================================
# Step 4 - Build a feature matrix (X) and label vector (y) for a split
# =============================================================================
def build_dataset(img_dir, lbl_dir, max_images, split_name, samples=None):
    """Return (X, y) for one split by cropping patches and extracting LBPH.

    If `samples` (a dict) is provided, a few example face / non-face patches
    are stored in it for later visualisation.
    """
    image_files = list_images(img_dir)
    if max_images is not None and len(image_files) > max_images:
        image_files = random.sample(image_files, max_images)

    X, y = [], []
    n_pos, n_neg = 0, 0

    for idx, img_path in enumerate(image_files, 1):
        img = cv2.imread(img_path, cv2.IMREAD_GRAYSCALE)
        if img is None:
            continue
        h, w = img.shape[:2]

        stem = os.path.splitext(os.path.basename(img_path))[0]
        label_path = os.path.join(lbl_dir, stem + ".txt")
        face_boxes = read_yolo_labels(label_path, w, h)

        # --- positive patches (faces) ---
        for box in face_boxes[:MAX_FACES_PER_IMAGE]:
            patch = crop_patch(img, box)
            if patch is not None:
                X.append(extract_lbph(patch))
                y.append(1)
                n_pos += 1
                if samples is not None and len(samples["face"]) < 5:
                    samples["face"].append(patch)

        # --- negative patches (background) ---
        for box in sample_negative_boxes(face_boxes, w, h, NEG_PER_IMAGE):
            patch = crop_patch(img, box)
            if patch is not None:
                X.append(extract_lbph(patch))
                y.append(0)
                n_neg += 1
                if samples is not None and len(samples["nonface"]) < 5:
                    samples["nonface"].append(patch)

        if idx % 100 == 0:
            print(f"    [{split_name}] processed {idx}/{len(image_files)} images "
                  f"-> {n_pos} faces, {n_neg} non-faces")

    print(f"    [{split_name}] DONE: {n_pos} face + {n_neg} non-face patches "
          f"from {len(image_files)} images")
    return np.asarray(X, dtype=np.float32), np.asarray(y, dtype=np.int32)


# =============================================================================
# Step 6b - Visualisations (charts saved as PNG files)
# =============================================================================
def plot_sample_patches(samples, out_path):
    """Save a grid of example face / non-face patches and their LBP image."""
    faces = samples["face"][:5]
    nonfaces = samples["nonface"][:5]
    n = max(len(faces), len(nonfaces))
    if n == 0:
        return None
    fig, axes = plt.subplots(2, n, figsize=(2 * n, 4.2))
    if n == 1:
        axes = axes.reshape(2, 1)
    for col in range(n):
        for row, (group, title) in enumerate([(faces, "FACE"),
                                               (nonfaces, "NON-FACE")]):
            ax = axes[row, col]
            ax.axis("off")
            if col < len(group):
                ax.imshow(group[col], cmap="gray")
                if col == 0:
                    ax.set_title(title, fontsize=11, loc="left")
    fig.suptitle("Sample training patches", fontsize=13)
    fig.tight_layout()
    fig.savefig(out_path, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_lbp_example(sample_patch, out_path):
    """Save input patch, its LBP image and the LBP histogram side by side."""
    if sample_patch is None:
        return None
    lbp = local_binary_pattern(sample_patch, LBP_POINTS, LBP_RADIUS, LBP_METHOD)
    n_bins = LBP_POINTS + 2
    fig, axes = plt.subplots(1, 3, figsize=(11, 3.4))
    axes[0].imshow(sample_patch, cmap="gray"); axes[0].set_title("Input patch")
    axes[0].axis("off")
    axes[1].imshow(lbp, cmap="gray"); axes[1].set_title("LBP image")
    axes[1].axis("off")
    axes[2].hist(lbp.ravel(), bins=n_bins, range=(0, n_bins),
                 color="#3b6ea5", edgecolor="black")
    axes[2].set_title("LBP histogram"); axes[2].set_xlabel("LBP code")
    axes[2].set_ylabel("count")
    fig.suptitle("Computation of LBP / LBPH feature", fontsize=13)
    fig.tight_layout()
    fig.savefig(out_path, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_confusion_matrix(cm, out_path):
    """Save the confusion matrix as an annotated heatmap."""
    fig, ax = plt.subplots(figsize=(4.6, 4.2))
    im = ax.imshow(cm, cmap="Blues")
    labels = ["non-face", "face"]
    ax.set_xticks([0, 1]); ax.set_xticklabels(labels)
    ax.set_yticks([0, 1]); ax.set_yticklabels(labels)
    ax.set_xlabel("Predicted"); ax.set_ylabel("True")
    ax.set_title("Confusion matrix (validation set)")
    thresh = cm.max() / 2.0 if cm.max() > 0 else 0.5
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(cm[i, j]), ha="center", va="center",
                    color="white" if cm[i, j] > thresh else "black",
                    fontsize=14)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(out_path, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_metrics_bar(metrics, out_path):
    """Save a bar chart of accuracy / precision / recall / F1."""
    names = list(metrics.keys())
    values = list(metrics.values())
    fig, ax = plt.subplots(figsize=(5.2, 3.6))
    bars = ax.bar(names, values, color=["#4c72b0", "#55a868",
                                        "#c44e52", "#8172b3"])
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("score")
    ax.set_title("Validation metrics")
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, v + 0.02, f"{v:.3f}",
                ha="center", fontsize=10)
    fig.tight_layout()
    fig.savefig(out_path, dpi=120, bbox_inches="tight")
    plt.close(fig)
    return out_path


# =============================================================================
# Step 7b - Build a Word (.docx) report
# =============================================================================
def _report_texts(lang, feat_len):
    """Return all localised strings for the report in the chosen language."""
    if lang == "vi":
        return {
            "title": "Nhận dạng khuôn mặt dùng đặc trưng LBPH + SVM",
            "info": [("Họ và tên", STUDENT_NAME), ("MSHV", STUDENT_ID),
                     ("Môn", COURSE), ("Bài tập ngày", ASSIGNMENT_DATE)],
            "h_obj": "1. Mục tiêu",
            "obj": (
                "Xây dựng chương trình nhận dạng khuôn mặt sử dụng đặc trưng "
                "LBPH (Local Binary Pattern Histogram) kết hợp với bộ phân lớp "
                "SVM (libsvm). Bộ dữ liệu face-detection-dataset trên Kaggle "
                "cung cấp các hộp bao (bounding box) khuôn mặt, nên bài toán "
                "được mô hình hoá thành bài toán phân lớp nhị phân: phân biệt "
                "vùng KHUÔN MẶT (cắt từ các hộp bao) và vùng KHÔNG PHẢI KHUÔN "
                "MẶT (vùng nền)."),
            "h_method": "2. Phương pháp",
            "method": [
                "Bước 1 - Đọc ảnh và nhãn định dạng YOLO của tập train/val.",
                "Bước 2 - Cắt các vùng khuôn mặt từ hộp bao (mẫu dương) và các "
                "vùng nền ngẫu nhiên không trùng khuôn mặt (mẫu âm).",
                f"Bước 3 - Đưa mỗi vùng ảnh về kích thước "
                f"{PATCH_SIZE[0]}x{PATCH_SIZE[1]} và trích xuất đặc trưng LBPH: "
                f"LBP uniform (P={LBP_POINTS}, R={LBP_RADIUS}) tính trên lưới "
                f"{GRID[0]}x{GRID[1]} ô, mỗi ô một histogram đã chuẩn hoá, ghép "
                f"lại thành vector {feat_len} chiều.",
                "Bước 4 - Chuẩn hoá đặc trưng (trung bình 0, phương sai 1).",
                "Bước 5 - Huấn luyện SVM (libsvm C-SVC, nhân RBF) trên tập train.",
                "Bước 6 - Đánh giá trên tập val bằng độ chính xác (accuracy), "
                "precision, recall, F1-score và ma trận nhầm lẫn.",
            ],
            "cap_lbp": "Minh hoạ quá trình tính LBP / LBPH:",
            "cap_samples": "Một số vùng ảnh mẫu được trích xuất để huấn luyện:",
            "h_setup": "3. Thiết lập thực nghiệm",
            "setup": [
                ("Số vùng ảnh train", str), ("Số vùng ảnh val", str),
                ("Số chiều đặc trưng", str), ("Số vector hỗ trợ", str),
                ("LBP", f"uniform, P={LBP_POINTS}, R={LBP_RADIUS}"),
                ("Lưới ô", f"{GRID[0]} x {GRID[1]} ô"),
                ("Bộ phân lớp", "SVM (libsvm), nhân RBF, C=10, gamma=scale"),
            ],
            "h_res": "4. Kết quả và đánh giá",
            "col_metric": "Chỉ số", "col_value": "Giá trị",
            "cm_text": (
                "Ma trận nhầm lẫn: thực tế không-phải-mặt được dự đoán là "
                "[không-mặt={a}, mặt={b}]; thực tế khuôn-mặt được dự đoán là "
                "[không-mặt={c}, mặt={d}]."),
            "h_ana": "5. Phân tích",
            "ana": [
                "LBPH mã hoá kết cấu (texture) cục bộ và ít nhạy với thay đổi "
                "độ sáng đơn điệu, nhờ đó ngay cả một SVM đơn giản cũng phân "
                "biệt khá tốt kết cấu khuôn mặt với vùng nền.",
                "Nguồn lỗi chính là các vùng nền có kết cấu giống mặt (vùng "
                "trơn / giống màu da) và các khuôn mặt quá nhỏ hoặc bị mờ, mất "
                "cấu trúc LBP đặc trưng sau khi co giãn kích thước.",
                "Hướng cải tiến: thêm nhiều mẫu âm khó hơn, tinh chỉnh C và "
                "gamma của SVM bằng kiểm chứng chéo (cross-validation), dùng "
                "lưới LBP mịn hơn, hoặc kết hợp LBPH với đặc trưng HOG.",
            ],
        }
    # default: English
    return {
        "title": "Face Recognition using LBPH features + SVM",
        "info": [("Họ và tên", STUDENT_NAME), ("MSHV", STUDENT_ID),
                 ("Môn", COURSE), ("Bài tập ngày", ASSIGNMENT_DATE)],
        "h_obj": "1. Objective",
        "obj": (
            "Build a face recognition program that uses the LBPH (Local Binary "
            "Pattern Histogram) feature descriptor together with an SVM "
            "classifier (libsvm). The Kaggle face-detection-dataset provides "
            "face bounding boxes, so the task is modelled as a binary "
            "classifier separating FACE patches (cropped from the boxes) from "
            "NON-FACE background patches."),
        "h_method": "2. Method",
        "method": [
            "Step 1 - Read images and YOLO-format labels of the train/val splits.",
            "Step 2 - Crop face patches from the bounding boxes (positive "
            "samples) and random background patches not overlapping faces "
            "(negative samples).",
            f"Step 3 - Resize every patch to {PATCH_SIZE[0]}x{PATCH_SIZE[1]} and "
            f"extract the LBPH descriptor: uniform LBP (P={LBP_POINTS}, "
            f"R={LBP_RADIUS}) computed over a {GRID[0]}x{GRID[1]} grid of cells, "
            f"one normalised histogram per cell, concatenated into a "
            f"{feat_len}-dim vector.",
            "Step 4 - Standardise the features (zero mean, unit variance).",
            "Step 5 - Train an SVM (libsvm C-SVC, RBF kernel) on the train split.",
            "Step 6 - Evaluate on the val split with accuracy, precision, "
            "recall, F1-score and the confusion matrix.",
        ],
        "cap_lbp": "Illustration of the LBP / LBPH computation:",
        "cap_samples": "Examples of the extracted training patches:",
        "h_setup": "3. Experimental setup",
        "setup": [
            ("Train patches", str), ("Validation patches", str),
            ("Feature dimension", str), ("Support vectors", str),
            ("LBP", f"uniform, P={LBP_POINTS}, R={LBP_RADIUS}"),
            ("Grid", f"{GRID[0]} x {GRID[1]} cells"),
            ("Classifier", "SVM (libsvm), RBF kernel, C=10, gamma=scale"),
        ],
        "h_res": "4. Results and evaluation",
        "col_metric": "Metric", "col_value": "Value",
        "cm_text": (
            "Confusion matrix: true non-face classified as "
            "[non-face={a}, face={b}]; true face classified as "
            "[non-face={c}, face={d}]."),
        "h_ana": "5. Analysis",
        "ana": [
            "LBPH encodes local texture and is robust to monotonic illumination "
            "changes, which lets even a simple SVM separate facial texture from "
            "background reasonably well.",
            "The main error source is background patches whose texture resembles "
            "a face (smooth / skin-like regions) and very small or blurry faces "
            "that lose discriminative LBP structure after resizing.",
            "Possible improvements: add more and harder negative samples, tune "
            "the SVM C and gamma via cross-validation, use a finer LBP grid, or "
            "combine LBPH with HOG features.",
        ],
    }


def build_word_report(metrics, cm, n_train, n_val, feat_len, n_sv,
                      chart_paths, out_path, lang="en"):
    """Generate the assignment report as a .docx file with embedded charts.

    `lang` selects the language of the body text: "en" or "vi".
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  (python-docx not installed -> skipping Word report)")
        return None

    t = _report_texts(lang, feat_len)
    # numeric values that fill the `str` placeholders of the setup table
    setup_values = [str(n_train), str(n_val), str(feat_len), str(n_sv)]

    doc = Document()

    # --- Title + student information block ---
    title = doc.add_heading(t["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    for label, value in t["info"]:
        run = info.add_run(f"{label}: ")
        run.bold = True
        info.add_run(f"{value}\n")

    # --- 1. Objective ---
    doc.add_heading(t["h_obj"], level=1)
    doc.add_paragraph(t["obj"])

    # --- 2. Method ---
    doc.add_heading(t["h_method"], level=1)
    for line in t["method"]:
        doc.add_paragraph(line, style="List Bullet")

    if chart_paths.get("lbp") and os.path.exists(chart_paths["lbp"]):
        doc.add_paragraph(t["cap_lbp"])
        doc.add_picture(chart_paths["lbp"], width=Inches(6.0))
    if chart_paths.get("samples") and os.path.exists(chart_paths["samples"]):
        doc.add_paragraph(t["cap_samples"])
        doc.add_picture(chart_paths["samples"], width=Inches(5.5))

    # --- 3. Experimental setup ---
    doc.add_heading(t["h_setup"], level=1)
    setup = doc.add_table(rows=0, cols=2)
    setup.style = "Light Grid Accent 1"
    vi = 0
    for key, val in t["setup"]:
        if val is str:               # placeholder -> next numeric value
            val = setup_values[vi]
            vi += 1
        row = setup.add_row().cells
        row[0].text = key
        row[1].text = val

    # --- 4. Results ---
    doc.add_heading(t["h_res"], level=1)
    res = doc.add_table(rows=1, cols=2)
    res.style = "Light Grid Accent 1"
    res.rows[0].cells[0].text = t["col_metric"]
    res.rows[0].cells[1].text = t["col_value"]
    for k, v in metrics.items():
        row = res.add_row().cells
        row[0].text = k
        row[1].text = f"{v:.4f}"

    if chart_paths.get("metrics") and os.path.exists(chart_paths["metrics"]):
        doc.add_picture(chart_paths["metrics"], width=Inches(4.6))
    if chart_paths.get("cm") and os.path.exists(chart_paths["cm"]):
        doc.add_picture(chart_paths["cm"], width=Inches(4.2))

    doc.add_paragraph(t["cm_text"].format(a=cm[0, 0], b=cm[0, 1],
                                          c=cm[1, 0], d=cm[1, 1]))

    # --- 5. Analysis ---
    doc.add_heading(t["h_ana"], level=1)
    for line in t["ana"]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(out_path)
    return out_path


# =============================================================================
# Step 5 - Train the SVM and Step 6 - Evaluate it
# =============================================================================
def main():
    print("=" * 70)
    print("Face recognition with LBPH features + SVM (libsvm)")
    print("=" * 70)

    # --- Step 1: locate data ---
    print("\n[Step 1] Locating dataset splits ...")
    splits = find_split_dirs(DATASET_DIR)
    if "train" not in splits or "val" not in splits:
        print(f"  ERROR: could not find train/val folders under {DATASET_DIR}")
        print("  Expected e.g. dataset/images/train + dataset/labels/train")
        print(f"  Found: {splits}")
        return
    print(f"  train images: {splits['train'][0]}")
    print(f"  val   images: {splits['val'][0]}")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- Steps 2-4: build feature matrices ---
    print("\n[Step 2-4] Extracting LBPH features for the TRAIN split ...")
    samples = {"face": [], "nonface": []}
    X_train, y_train = build_dataset(*splits["train"], MAX_IMAGES_TRAIN, "train",
                                     samples=samples)

    print("\n[Step 2-4] Extracting LBPH features for the VAL split ...")
    X_val, y_val = build_dataset(*splits["val"], MAX_IMAGES_VAL, "val")

    if len(X_train) == 0 or len(X_val) == 0:
        print("  ERROR: no patches were extracted. Is the dataset downloaded?")
        return

    print(f"\n  Feature vector length : {X_train.shape[1]}")
    print(f"  Train samples         : {X_train.shape[0]}")
    print(f"  Val   samples         : {X_val.shape[0]}")

    # --- Step 5: standardise features + train the SVM ---
    print("\n[Step 5] Training the SVM classifier (libsvm, RBF kernel) ...")
    scaler = StandardScaler().fit(X_train)
    X_train_s = scaler.transform(X_train)
    X_val_s = scaler.transform(X_val)

    # SVC == scikit-learn's wrapper around libsvm (C-SVC, RBF kernel).
    clf = SVC(kernel="rbf", C=10.0, gamma="scale",
              class_weight="balanced", random_state=RANDOM_SEED)
    clf.fit(X_train_s, y_train)
    print(f"  Support vectors used: {int(clf.support_vectors_.shape[0])}")

    # --- Step 6: evaluate on the validation split ---
    print("\n[Step 6] Evaluating on the VAL split ...")
    y_pred = clf.predict(X_val_s)

    acc = accuracy_score(y_val, y_pred)
    prec = precision_score(y_val, y_pred, zero_division=0)
    rec = recall_score(y_val, y_pred, zero_division=0)
    f1 = f1_score(y_val, y_pred, zero_division=0)
    cm = confusion_matrix(y_val, y_pred)

    print("\n" + "-" * 70)
    print("RESULTS ON VALIDATION SET")
    print("-" * 70)
    print(f"  Accuracy : {acc:.4f}")
    print(f"  Precision: {prec:.4f}  (of patches predicted FACE, how many are)")
    print(f"  Recall   : {rec:.4f}  (of real faces, how many were found)")
    print(f"  F1-score : {f1:.4f}")
    print("\n  Confusion matrix (rows = true, cols = predicted):")
    print("                 pred non-face   pred face")
    print(f"    true non-face    {cm[0, 0]:>8d}    {cm[0, 1]:>8d}")
    print(f"    true face        {cm[1, 0]:>8d}    {cm[1, 1]:>8d}")
    print("\n  Per-class report:")
    print(classification_report(y_val, y_pred,
                                target_names=["non-face", "face"],
                                zero_division=0))

    # --- Step 6b: visual charts ---
    print("\n[Step 6b] Saving visual charts to outputs/ ...")
    metrics = {"Accuracy": acc, "Precision": prec, "Recall": rec, "F1-score": f1}
    chart_paths = {
        "samples": plot_sample_patches(
            samples, os.path.join(OUTPUT_DIR, "sample_patches.png")),
        "lbp": plot_lbp_example(
            samples["face"][0] if samples["face"] else None,
            os.path.join(OUTPUT_DIR, "lbp_example.png")),
        "cm": plot_confusion_matrix(
            cm, os.path.join(OUTPUT_DIR, "confusion_matrix.png")),
        "metrics": plot_metrics_bar(
            metrics, os.path.join(OUTPUT_DIR, "metrics_bar.png")),
    }
    for name, path in chart_paths.items():
        if path:
            print(f"    saved {os.path.relpath(path, SCRIPT_DIR)}")

    # --- Step 7: short written analysis ---
    print("\n" + "-" * 70)
    print("ANALYSIS")
    print("-" * 70)
    print(
        "  * LBPH encodes local texture and is robust to monotonic illumination\n"
        "    changes, which is why it separates facial texture from background\n"
        "    fairly well even with a simple SVM.\n"
        "  * Main error source: background patches whose texture resembles a\n"
        "    face (skin-coloured / smooth regions) and very small or blurry\n"
        "    faces that lose discriminative LBP structure once resized.\n"
        "  * To improve: add more (and harder) negative samples, tune SVM C and\n"
        "    gamma via cross-validation, use a finer LBP grid, or combine LBPH\n"
        "    with HOG features.")

    # --- Step 7b: Word reports (English + Vietnamese) ---
    print("\n[Step 7b] Building the Word reports (EN + VI) ...")
    for lang, fname in [
        ("en", "Report_FaceRecognition_LBPH_SVM_EN.docx"),
        ("vi", "BaoCao_NhanDangKhuonMat_LBPH_SVM_VI.docx"),
    ]:
        report_path = build_word_report(
            metrics, cm, X_train.shape[0], X_val.shape[0], X_train.shape[1],
            int(clf.support_vectors_.shape[0]), chart_paths,
            os.path.join(OUTPUT_DIR, fname), lang=lang)
        if report_path:
            print(f"    saved {os.path.relpath(report_path, SCRIPT_DIR)}")
    print("=" * 70)


if __name__ == "__main__":
    main()
